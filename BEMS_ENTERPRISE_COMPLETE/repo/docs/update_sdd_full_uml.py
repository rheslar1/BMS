from docx import Document
from docx.shared import Inches
import subprocess
import os

base_dir = os.path.dirname(os.path.abspath(__file__))

uml_defs = {
    'architecture': '''digraph Architecture {
  rankdir=LR;
  node [shape=box, style=filled, fillcolor="#e8f4ff", color="#365f91", fontname="Arial"];
  BrowserNode [label="Browser\nReact UI\nSSE / HTTP"];
  APINode [label="Node.js API\nHTTP/JSON / SSE"];
  CoreNode [label="C++ Edge Core\nAI + Analytics + BACnet"];
  DeviceNode [label="BACnet/IP Devices\nFreeRTOS Field Devices"];
  BrowserNode -> APINode [label="HTTP/JSON", color="#365f91"];
  APINode -> CoreNode [label="gRPC", color="#365f91"];
  CoreNode -> DeviceNode [label="BACnet/IP", color="#365f91"];
}''',
    'layers': '''digraph Layers {
  rankdir=TB;
  node [shape=record, style=filled, fillcolor="#f5f7fb", color="#365f91", fontname="Arial"];
  Presentation [label="{Presentation Layer|React UI\nSSE real-time updates}"];
  Application [label="{Application Layer|Node.js API\nHTTP/JSON + SSE}"];
  ServiceLayer [label="{Service Layer|gRPC services\nNode.js ↔ C++ Edge Core}"];
  CoreLayer [label="{Core Layer|C++ edge core\nControl, AI, Analytics, BACnet}"];
  FieldLayer [label="{Field Layer|FreeRTOS BACnet/IP devices\nEEPROM config}"];
  Presentation -> Application -> ServiceLayer -> CoreLayer -> FieldLayer;
}''',
    'data_model': '''digraph DataModel {
  rankdir=LR;
  node [shape=record, style=filled, fillcolor="#f9f9f9", color="#365f91", fontname="Arial"];
  BuildingEntity [label="{Building|building_id: int\nname: string\naddress: string\ndescription: text}"];
  ZoneEntity [label="{Zone|zone_id: int\nbuilding_id: int\nname: string\ndescription: text}"];
  DeviceEntity [label="{Device|device_id: int\nzone_id: int\nname: string\ntype: string\nbacnet_instance: int\nobject_type: string\nvendor: string\nmodel: string\nip_address: string\npresent_value: double\nunits: string\nstatus: string\ndescription: text}"];
  AlarmEntity [label="{Alarm|id: int\nmessage: text}"];
  BuildingEntity -> ZoneEntity [label="1..*", color="#365f91"];
  ZoneEntity -> DeviceEntity [label="1..*", color="#365f91"];
}''',
    'sequence': '''digraph Sequence {
  rankdir=LR;
  node [shape=box, style=filled, fillcolor="#eef6ff", color="#365f91", fontname="Arial"];
  BrowserNode [label="Browser UI"];
  APINode [label="Node.js API"];
  CoreNode [label="C++ Edge Core"];
  DeviceNode [label="BACnet/IP Device"];
  BrowserNode -> APINode [label="Request data / command"];
  APINode -> CoreNode [label="gRPC call"];
  CoreNode -> DeviceNode [label="BACnet/IP read/write"];
  DeviceNode -> CoreNode [label="BACnet/IP response"];
  CoreNode -> APINode [label="gRPC response"];
  APINode -> BrowserNode [label="HTTP/JSON + SSE update"];
}''',
}

image_paths = {}
for name, dot in uml_defs.items():
    png_path = os.path.join(base_dir, f"uml_{name}.png")
    proc = subprocess.Popen(['dot', '-Tpng'], stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    png_data, err = proc.communicate(dot.encode('utf-8'))
    if proc.returncode != 0:
        raise RuntimeError(f"Graphviz failed for {name}: {err.decode('utf-8')}")
    with open(png_path, 'wb') as f:
        f.write(png_data)
    image_paths[name] = png_path

sdd_path = os.path.join(base_dir, 'SDD.docx')
doc = Document(sdd_path)

if len(doc.paragraphs) > 0:
    doc.add_page_break()

doc.add_heading('UML Diagrams', level=1)
doc.add_paragraph('This section presents UML diagrams for the system architecture, layered deployment, data model, and platform sequence flow.')

doc.add_heading('System Architecture', level=2)
doc.add_paragraph('Architecture diagram showing browser UI, Node.js API, C++ edge core, and BACnet/IP field devices.')
doc.add_picture(image_paths['architecture'], width=Inches(6.5))

doc.add_heading('Layered Architecture', level=2)
doc.add_paragraph('Layered architecture from presentation down to field devices.')
doc.add_picture(image_paths['layers'], width=Inches(6.5))

doc.add_heading('Data Model', level=2)
doc.add_paragraph('Class diagram showing Building, Zone, Device, and Alarm entities.')
doc.add_picture(image_paths['data_model'], width=Inches(6.5))

doc.add_heading('Message Flow', level=2)
doc.add_paragraph('Sequence diagram showing data flow from browser to device and back.')
doc.add_picture(image_paths['sequence'], width=Inches(6.5))

doc.save(sdd_path)
print('Inserted UML diagrams into', sdd_path)

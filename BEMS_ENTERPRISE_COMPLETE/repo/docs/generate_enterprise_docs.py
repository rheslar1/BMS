from pathlib import Path

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parent


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_image_if_present(doc, filename, caption):
    path = BASE_DIR / filename
    if path.exists():
        doc.add_paragraph(caption)
        doc.add_picture(str(path), width=Inches(6.4))


def build_sdd():
    doc = Document()
    doc.add_heading("BMS AI Edge Platform - Software Design Document", 0)

    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "This document defines the software design for a commercial-grade Building Management "
        "and Energy Management System. The platform combines a React operator UI, Node.js REST/SSE "
        "API, MySQL storage, Python AI optimizer, RabbitMQ-orchestrated C++ edge core, and BACnet/IP field integration."
    )
    doc.add_paragraph(
        "The platform unifies HVAC control, lighting supervision, power monitoring, and energy "
        "optimization in one BMS/BEMS GUI and API surface."
    )
    doc.add_paragraph(
        "The design is a modern, scalable implementation of the traditional BACnet/BMS architecture. "
        "BACnet remains the device and field integration layer, while analytics, AI optimization, "
        "remote management, and enterprise administration are layered above it."
    )
    doc.add_paragraph(
        "Internally, BACnet concepts drive discovery, object modeling, telemetry, writeback, schedules, "
        "alarms, and controller-to-server communication. The enterprise layer adds operator UI, analytics, "
        "cloud-style API integration, health monitoring, remote management, and scalability from a single "
        "building to distributed campus deployments."
    )
    doc.add_paragraph(
        "The architecture is best suited for smart buildings, energy-focused facilities, mixed HVAC and "
        "power systems, airports, hospitals, large campuses, and enterprise-scale energy + IoT integration."
    )
    add_bullets(
        doc,
        [
            "Runs dashboards for real buildings and BACnet-connected HVAC equipment.",
            "Stores building hierarchy, telemetry, alarms, analytics, users, roles, sessions, and AI history in MySQL.",
            "Uses Server-Sent Events for browser live updates and does not use WebSockets.",
            "Uses gRPC for Node-to-AI optimization and RabbitMQ AMQP for Node-to-C++ edge command orchestration.",
            "Deploys through Ubuntu 22.04 Docker containers and has an i.MX93 Yocto packaging path.",
        ],
    )

    doc.add_heading("2. Complete System Architecture", level=1)
    doc.add_paragraph("Runtime chain:")
    add_numbered(
        doc,
        [
            "React UI sends HTTP/JSON commands and receives SSE telemetry/alarm updates.",
            "Node.js API enforces authentication, RBAC, tenant context, orchestration, and persistence.",
            "MySQL stores configuration, alarms, analytics, optimization history, and RL Q-values.",
            "Node.js calls the Python AI service over gRPC for whole-building optimization.",
            "Node.js queues C++ edge-core discovery, read, write, COV, OTA, and forecast commands through RabbitMQ AMQP.",
            "The C++ edge core communicates with BACnet/IP networks over UDP/BVLC.",
            "HVAC devices expose BACnet objects for telemetry, status, and command writeback.",
        ],
    )
    add_image_if_present(doc, "uml_architecture.png", "Architecture diagram")
    add_image_if_present(doc, "uml_layers.png", "Layered deployment diagram")
    add_image_if_present(doc, "uml_sequence.png", "Sequence diagram")

    doc.add_heading("3. Major Components", level=1)
    add_bullets(
        doc,
        [
            "ui/: React dashboard with login, admin, user maintenance, alarms, schedules, charts, digital twin, floorplan editor, provisioning, and AI control panels.",
            "node-api/: Express API with REST endpoints, SSE streams, RabbitMQ edge client, AI gRPC client, authentication, authorization, migrations, trend logging, holiday/special-event scheduling, Smart Grid AI, watchdog, and remote management.",
            "ai-service/: Python optimizer exposing AiOptimizationService over gRPC and HTTP health/fallback endpoints.",
            "edge-core/: C++ BACnet runtime with RabbitMQ edge command orchestration, BACnet server/device object database, Who-Is, ReadProperty, ReadPropertyMultiple, WriteProperty, SubscribeCOV, and COV notification support.",
            "database/: Canonical MySQL schema for enterprise SaaS, telemetry, alarms, analytics, RL, FDD, and maintenance data.",
            "docker/: Ubuntu-container Docker Compose stack for API, UI, AI service, edge core, and MySQL.",
            "yocto/: i.MX93 deployment layer and systemd packaging path.",
        ],
    )

    doc.add_heading("4. Service Contracts", level=1)
    add_bullets(
        doc,
        [
            "RabbitMQ topic bems.edge.commands carries edge command types bacnet.discover_devices, bacnet.read_property, bacnet.read_property_multiple, bacnet.write_property, bacnet.subscribe_cov, edge.energy_forecast, and nrf52840.ota_update.",
            "proto/ai_service.proto defines AiOptimizationService: Health, Optimize, and Feedback.",
            "node-api/edgeClient.js and node-api/aiClient.js are the client-side adapters.",
            "edge-core/src/main.cpp runs the C++ BACnet runtime for RabbitMQ command orchestration and BACnet/IP field integration.",
            "ai-service/app.py is the Python AiOptimizationService implementation.",
        ],
    )

    doc.add_heading("5. BACnet/IP Design", level=1)
    add_bullets(
        doc,
        [
            "C++ BACnet integration uses UDP socket initialization and BACnet/IP BVLC framing.",
            "Discovery follows Who-Is and I-Am patterns for BACnet device lifecycle support.",
            "Read operations use confirmed ReadProperty for present-value telemetry.",
            "Write operations use confirmed WriteProperty for present-value command output.",
            "COV setup uses confirmed SubscribeCOV for BACnet object change monitoring.",
            "BACnet object examples include temperature sensor as Analog Input, fan command as Binary Output, damper command as Analog Output, power meter as Analog Input/Value, and schedule as BACnet Schedule Object.",
            "A built-in simulator exposes BACnet instances 101, 102, 103, 201, 250, 301, and 302 through the same discovery/read/write boundary when BACNET_SIMULATOR_ENABLED=true.",
            "Docker Compose enables the simulator for demos; embedded systemd deployment defaults it off for real BACnet/IP networks.",
            "Safe writeback uses a strategy object for absolute/delta modes, clamping, verification, and rollback.",
        ],
    )
    doc.add_paragraph(
        "Field-level connected products include sensors, actuators, VFD drives, power meters, "
        "smart breakers, SmartX-style controllers, PLCs, and programmable automation controllers. "
        "BACnet MS/TP, Modbus RTU, CAN bus, and Zigbee/wireless networks are treated as adapter or gateway "
        "integration paths into the building model."
    )

    doc.add_heading("6. HVAC and Control Modeling", level=1)
    add_bullets(
        doc,
        [
            "Building model: Building -> Floor -> Room -> Zone -> Device -> BACnet object metadata.",
            "Thermal features include outdoor temperature, solar radiation, wind, envelope quality, occupancy, internal loads, SAT, VAV flows, zone temperature, runtime history, and warm-up/cool-down curves.",
            "Autonomous profiles are Conservative, Normal, and Aggressive.",
            "The AI control loop keeps occupied people comfortable, minimizes energy, and avoids overload peaks.",
            "The digital twin mirrors zones, floors, rooms, devices, status, values, overlays, provisioning state, and summaries in real time.",
            "Schedules support daily, monthly, yearly, holiday, and special-event override behavior.",
            "Trend logging persists point samples for graphics, reports, diagnostics, FDD, and AI analytics.",
            "Smart Grid AI coordinates HVAC, lighting, power monitoring, storage, price signals, and demand response while respecting fire and security priorities.",
            "Predictive simulation evaluates proposed setpoint actions before optional writeback.",
            "The airflow graph model is structured for a future graph neural network upgrade.",
        ],
    )

    doc.add_heading("7. Reinforcement Learning Model", level=1)
    add_bullets(
        doc,
        [
            "State: current building condition across zones, devices, weather, pricing, occupancy, demand response, and peak-load context.",
            "Action: control decisions such as setpoint adjustment, schedule profile, ventilation bias, or hold action.",
            "Reward: performance score balancing comfort protection, energy reduction, cost reduction, and peak avoidance.",
            "Policy: learned action strategy persisted as zone/action Q-values in MySQL.",
            "History: optimization runs are stored for audit, analytics, explainability, and retraining.",
        ],
    )

    doc.add_heading("8. Security and Administration", level=1)
    add_bullets(
        doc,
        [
            "Users authenticate with salted scrypt password hashes.",
            "UI sessions use X-Session-Token.",
            "API integrations use session tokens issued through the login endpoint.",
            "Role permissions protect mutating endpoints and administrative actions.",
            "Audit events record security and administrative operations.",
        ],
    )

    doc.add_heading("9. UML and Data Model", level=1)
    add_image_if_present(doc, "uml_data_model.png", "Data model diagram")
    doc.add_paragraph(
        "Primary tables include organizations, sites, users, roles, sessions, audit events, "
        "buildings, floors, rooms, zones, devices, alarms, alarm logs, trend logs, schedules, holiday schedules, special events, analytics events, RL Q-values, optimization history, "
        "FDD findings, and maintenance tickets."
    )

    doc.add_heading("10. Production Boundaries", level=1)
    doc.add_paragraph("EcoStruxure-style feature mapping:")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "This System"
    hdr[1].text = "Enterprise BMS Equivalent"
    hdr[2].text = "Notes"
    for left, right, notes in [
        ("React UI", "WebStation UI", "Browser operator graphics, alarms, schedules, device details, dashboards, and admin tools."),
        ("Device tree", "System Tree", "Building -> floor -> room -> zone -> device hierarchy with BACnet point metadata."),
        ("Floorplan / digital twin", "Graphics pages", "Equipment graphics, overlays, live values, status colors, and clickable controls."),
        ("Schedule editor", "BACnet Schedule Object", "Daily, monthly, yearly, holiday, and special-event scheduling with override precedence; device schedules persist on the BACnet device."),
        ("Alarm workflow", "Alarm Server", "Active alarms, acknowledge/clear workflow, logs, severity, status, and SSE updates."),
        ("Automation rules", "Script / Function Block", "Autonomous mode, AI control loop, deterministic decision engine, and future ML/rule expansion."),
        ("Bulk operations", "Multi-edit / bindings", "Provisioning, commissioning, schedule inheritance, maintenance scoping, and API-driven updates."),
        ("Analytics dashboard", "Enterprise analytics", "Energy KPIs, optimization history, RL state, FDD findings, weather/pricing context, and reports."),
        ("Remote API", "Enterprise integration API", "HTTP/JSON API, sessions, RBAC, audit events, and gRPC service contracts."),
    ]:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        row[2].text = notes

    add_bullets(
        doc,
        [
            "The architecture is commercial BEMS/SCADA class and comparable in scope to Siemens Desigo CC, Schneider EcoStruxure, and Niagara Framework.",
            "Site deployment still requires field certification, cybersecurity hardening, commissioning, and acceptance testing before controlling occupied buildings.",
            "Object-list heuristics, AHU/VAV auto-mapping, and FDD rules are implemented surfaces that should be tuned during commissioning for each vendor/device family.",
        ],
    )

    doc.save(BASE_DIR / "SDD.docx")


def build_sdp():
    doc = Document()
    doc.add_heading("BMS AI Edge Platform - Software Development Plan", 0)

    doc.add_heading("1. Lifecycle", level=1)
    add_numbered(
        doc,
        [
            "Architecture and requirements are maintained in docs/architecture.md and docs/api-surface.md.",
            "Development is organized by component: UI, Node API, AI service, edge core, database, Docker, and Yocto.",
            "Changes are verified with unit tests, C++ CTest, syntax checks, UI production build, and architecture verification.",
            "Docker Compose validates service integration before release.",
            "Yocto packaging carries the edge runtime to i.MX93-class embedded deployments.",
        ],
    )

    doc.add_heading("2. CI/CD", level=1)
    add_bullets(
        doc,
        [
            "CI builds the C++ edge core with CMake.",
            "CI runs C++ writeback/BACnet encoding tests through CTest.",
            "CI installs Node API dependencies, runs audit-compatible dependency installation, and executes Node tests.",
            "CI installs Python AI requirements and runs Python optimizer tests.",
            "CI installs UI dependencies and runs the Vite production build.",
            "CI runs architecture verification to enforce required files, API markers, RabbitMQ edge markers, SSE usage, and the no-WebSocket rule.",
        ],
    )

    doc.add_heading("3. Deployment", level=1)
    add_bullets(
        doc,
        [
            "Docker Compose services use Ubuntu 22.04 containers for api, ui, ai-service, and edge-core, plus MySQL 8 for db.",
            "API listens on port 3000.",
            "UI listens on port 5173.",
            "AI service exposes HTTP health on 8000 and gRPC on 50052.",
            "Edge core exposes BACnet/IP UDP on 47808 and receives edge commands through RabbitMQ AMQP.",
            "MySQL stores system configuration and operational history.",
            "Node API uses AI_GRPC_ENDPOINT=ai-service:50052 for AI and RABBITMQ_URL with EDGE_COMMAND_TRANSPORT=rabbitmq for edge orchestration.",
        ],
    )

    doc.add_heading("4. Environment and Configuration", level=1)
    add_bullets(
        doc,
        [
            "BEMS_REQUIRE_AUTH controls strict API authentication behavior.",
            "BEMS_MANAGEMENT_TOKEN protects remote management actions.",
            "BACNET_LOCAL_IP binds the edge BACnet interface.",
            "EDGE_COMMAND_TRANSPORT=rabbitmq selects RabbitMQ edge command orchestration.",
            "BACNET_SIMULATOR_ENABLED enables simulated BACnet devices for local demos and CI.",
            "TELEMETRY_STREAM_INTERVAL_MS controls SSE telemetry refresh cadence.",
        ],
    )

    doc.add_heading("5. Test Plan", level=1)
    add_bullets(
        doc,
        [
            "Node API: npm --prefix node-api test.",
            "Python AI: python -m unittest test_app.py in an environment with requirements installed.",
            "C++ edge: cmake configure/build and ctest.",
            "UI: npm --prefix ui run build.",
            "Architecture: ./scripts/verify_architecture.sh.",
            "Deployment smoke: docker-compose up --build -d and verify API health, UI, AI health, RabbitMQ event bus health through Node API, SSE telemetry, and provisioning discovery.",
            "BACnet simulator smoke: verify /api/provisioning/discover returns simulated devices when BACNET_SIMULATOR_ENABLED=true.",
        ],
    )

    doc.add_heading("6. Operations", level=1)
    add_bullets(
        doc,
        [
            "Watchdog endpoint reports API, database, AI, and edge service status.",
            "Remote management endpoints support restart, update intent, and watchdog execution.",
            "Alarms stream to the UI through SSE.",
            "Telemetry streams to charts and live feed through SSE.",
            "Maintenance tickets and FDD findings support operator workflows.",
            "Optimization history and RL Q-values support AI audit and retraining.",
        ],
    )

    doc.add_heading("7. Release Readiness", level=1)
    add_bullets(
        doc,
        [
            "Confirm no WebSocket implementation markers exist in project code.",
            "Confirm gRPC endpoints are configured for AI and edge services.",
            "Confirm database migrations run successfully on container startup.",
            "Confirm edge device BACnet UDP access and network segmentation are configured at the site.",
            "Complete cybersecurity review, backup/restore test, and site acceptance testing before production building control.",
        ],
    )

    doc.save(BASE_DIR / "SDP.docx")


if __name__ == "__main__":
    build_sdd()
    build_sdp()
    print("Generated SDD.docx and SDP.docx")
